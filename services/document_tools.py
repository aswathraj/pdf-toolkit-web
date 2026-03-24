from __future__ import annotations

import html
import io
import os
import re
import shutil
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Iterable, Optional, Sequence

import fitz
import pytesseract
from docx import Document
from openpyxl import Workbook, load_workbook
from openpyxl.utils import get_column_letter
from PIL import Image, ImageOps
from pypdf import PdfReader, PdfWriter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


class ProcessingError(Exception):
    pass


@dataclass
class ProcessingResult:
    files: list[Path]
    message: str = ""


ProgressCallback = Optional[Callable[[float, str], None]]


TOOL_DEFINITIONS = [
    {
        "key": "jpg-to-pdf",
        "title": "JPG to PDF",
        "summary": "Turn one or many JPG, JPEG, PNG, or WEBP images into a single PDF.",
        "multiple": True,
        "accept": ".jpg,.jpeg,.png,.webp,.bmp,.tif,.tiff",
        "button": "Convert Images",
        "badge": "Image",
    },
    {
        "key": "pdf-to-jpg",
        "title": "PDF to JPG",
        "summary": "Render each PDF page to a high-resolution JPG file.",
        "multiple": False,
        "accept": ".pdf",
        "button": "Export Pages",
        "badge": "Image",
    },
    {
        "key": "resize-image",
        "title": "Resize & Convert Image",
        "summary": "Resize by pixels, set output DPI, and convert images between JPG, PNG, WEBP, TIFF, or BMP.",
        "multiple": True,
        "accept": ".jpg,.jpeg,.png,.webp,.bmp,.tif,.tiff",
        "button": "Resize / Convert",
        "badge": "Image",
    },
    {
        "key": "word-to-pdf",
        "title": "Word to PDF",
        "summary": "Convert DOCX files into readable PDFs with headings, paragraphs, and tables.",
        "multiple": True,
        "accept": ".docx",
        "button": "Convert DOCX",
        "badge": "Office",
    },
    {
        "key": "pdf-to-word",
        "title": "PDF to Word",
        "summary": "Extract PDF text and tables into an editable DOCX document.",
        "multiple": True,
        "accept": ".pdf",
        "button": "Create DOCX",
        "badge": "Office",
    },
    {
        "key": "excel-to-pdf",
        "title": "Excel to PDF",
        "summary": "Convert XLSX sheets into paginated PDF tables.",
        "multiple": True,
        "accept": ".xlsx,.xlsm",
        "button": "Convert XLSX",
        "badge": "Office",
    },
    {
        "key": "pdf-to-excel",
        "title": "PDF to Excel",
        "summary": "Pull tables from PDFs into XLSX and fall back to line-by-line text when needed.",
        "multiple": True,
        "accept": ".pdf",
        "button": "Create XLSX",
        "badge": "Office",
    },
    {
        "key": "merge-pdf",
        "title": "Merge PDF",
        "summary": "Combine many PDFs into one document in the upload order.",
        "multiple": True,
        "accept": ".pdf",
        "button": "Merge Files",
        "badge": "PDF",
    },
    {
        "key": "split-pdf",
        "title": "Split PDF",
        "summary": "Split every page or extract custom page ranges from a PDF.",
        "multiple": False,
        "accept": ".pdf",
        "button": "Split PDF",
        "badge": "PDF",
    },
    {
        "key": "remove-watermark",
        "title": "Remove Watermark",
        "summary": "Strip annotations and optionally redact known watermark text from a PDF.",
        "multiple": True,
        "accept": ".pdf",
        "button": "Clean PDF",
        "badge": "Repair",
    },
    {
        "key": "ocr-pdf",
        "title": "OCR PDF",
        "summary": "Create searchable PDFs from scans or image files with Tesseract OCR.",
        "multiple": True,
        "accept": ".pdf,.jpg,.jpeg,.png,.webp,.bmp,.tif,.tiff",
        "button": "Run OCR",
        "badge": "OCR",
    },
]


IMAGE_SUFFIXES = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".tif", ".tiff"}
IMAGE_EXPORT_FORMATS = {
    "original": (None, None),
    "jpeg": ("JPEG", ".jpg"),
    "png": ("PNG", ".png"),
    "webp": ("WEBP", ".webp"),
    "tiff": ("TIFF", ".tiff"),
    "bmp": ("BMP", ".bmp"),
}


def clamp_progress(progress: float) -> float:
    return max(0.0, min(progress, 1.0))


def emit_progress(progress_callback: ProgressCallback, progress: float, detail: str) -> None:
    if progress_callback is not None:
        progress_callback(clamp_progress(progress), detail)


def overall_ratio(index: int, total: int, inner_ratio: float = 1.0) -> float:
    if total <= 0:
        return clamp_progress(inner_ratio)
    return clamp_progress(((index - 1) + clamp_progress(inner_ratio)) / total)


def count_pdf_pages(file_path: Path) -> int:
    require_suffix(file_path, {".pdf"}, "PDF")
    document = fitz.open(file_path)
    try:
        return document.page_count
    finally:
        document.close()


def run_tool(
    tool_key: str,
    files: list[Path],
    output_dir: Path,
    form_data,
    progress_callback: ProgressCallback = None,
) -> ProcessingResult:
    output_dir.mkdir(parents=True, exist_ok=True)

    handlers = {
        "jpg-to-pdf": lambda: images_to_pdf(files, output_dir, progress_callback=progress_callback),
        "pdf-to-jpg": lambda: pdf_to_images(files, output_dir, progress_callback=progress_callback),
        "resize-image": lambda: resize_and_convert_images(
            files,
            output_dir,
            width_px=form_data.get("width_px", ""),
            height_px=form_data.get("height_px", ""),
            dpi=form_data.get("dpi", ""),
            output_format=form_data.get("output_format", "original"),
            keep_aspect=form_data.get("keep_aspect", "on") in {"on", "1", "true", "yes"},
            progress_callback=progress_callback,
        ),
        "word-to-pdf": lambda: word_to_pdf(files, output_dir, progress_callback=progress_callback),
        "pdf-to-word": lambda: pdf_to_word(files, output_dir, progress_callback=progress_callback),
        "excel-to-pdf": lambda: excel_to_pdf(files, output_dir, progress_callback=progress_callback),
        "pdf-to-excel": lambda: pdf_to_excel(files, output_dir, progress_callback=progress_callback),
        "merge-pdf": lambda: merge_pdfs(files, output_dir, progress_callback=progress_callback),
        "split-pdf": lambda: split_pdf(
            files,
            output_dir,
            mode=form_data.get("split_mode", "all"),
            ranges=form_data.get("ranges", ""),
            progress_callback=progress_callback,
        ),
        "remove-watermark": lambda: remove_watermark(
            files,
            output_dir,
            watermark_text=form_data.get("watermark_text", "").strip(),
            progress_callback=progress_callback,
        ),
        "ocr-pdf": lambda: ocr_files(files, output_dir, progress_callback=progress_callback),
    }

    try:
        return handlers[tool_key]()
    except KeyError as exc:
        raise ProcessingError(f"Unknown tool: {tool_key}") from exc


def require_suffix(path: Path, allowed: Iterable[str], label: str) -> None:
    if path.suffix.lower() not in set(allowed):
        raise ProcessingError(f"{path.name} is not a supported {label} file.")


def clean_text(value) -> str:
    if value is None:
        return ""
    return str(value).strip()


def chunk_list(values: Sequence, size: int) -> list[Sequence]:
    return [values[index : index + size] for index in range(0, len(values), size)]


def parse_page_ranges(ranges: str, max_pages: int) -> list[tuple[int, int]]:
    cleaned = ranges.replace(" ", "")
    if not cleaned:
        raise ProcessingError("Enter page ranges like 1-3,5,8-10.")

    resolved: list[tuple[int, int]] = []
    for piece in cleaned.split(","):
        if not piece:
            continue
        if "-" in piece:
            start_text, end_text = piece.split("-", 1)
            if not start_text.isdigit() or not end_text.isdigit():
                raise ProcessingError("Page ranges must contain numbers only.")
            start = int(start_text)
            end = int(end_text)
        else:
            if not piece.isdigit():
                raise ProcessingError("Page ranges must contain numbers only.")
            start = end = int(piece)

        if start < 1 or end < 1 or start > max_pages or end > max_pages or start > end:
            raise ProcessingError(f"Page range {piece} is outside the PDF page count.")
        resolved.append((start, end))

    if not resolved:
        raise ProcessingError("No valid page ranges were found.")
    return resolved


def parse_optional_positive_int(value, label: str) -> int | None:
    text = clean_text(value)
    if not text:
        return None
    if not text.isdigit() or int(text) <= 0:
        raise ProcessingError(f"{label} must be a positive whole number.")
    return int(text)


def infer_source_image_format(file_path: Path) -> tuple[str, str]:
    suffix = file_path.suffix.lower()
    suffix_to_format = {
        ".jpg": ("JPEG", ".jpg"),
        ".jpeg": ("JPEG", ".jpg"),
        ".png": ("PNG", ".png"),
        ".webp": ("WEBP", ".webp"),
        ".bmp": ("BMP", ".bmp"),
        ".tif": ("TIFF", ".tiff"),
        ".tiff": ("TIFF", ".tiff"),
    }
    try:
        return suffix_to_format[suffix]
    except KeyError as exc:
        raise ProcessingError(f"{file_path.name} is not a supported image file.") from exc


def resolve_resize_dimensions(
    width_px: int | None,
    height_px: int | None,
    source_width: int,
    source_height: int,
    keep_aspect: bool,
) -> tuple[int, int]:
    if width_px is None and height_px is None:
        return source_width, source_height

    if not keep_aspect:
        if width_px is None or height_px is None:
            raise ProcessingError("Enter both width and height when aspect ratio is turned off.")
        return width_px, height_px

    if width_px is not None and height_px is not None:
        scale = min(width_px / source_width, height_px / source_height)
        return max(1, round(source_width * scale)), max(1, round(source_height * scale))

    if width_px is not None:
        scale = width_px / source_width
        return width_px, max(1, round(source_height * scale))

    scale = height_px / source_height
    return max(1, round(source_width * scale)), height_px


def save_image(image: Image.Image, destination: Path, image_format: str, dpi: int | None) -> None:
    save_kwargs = {}
    if dpi is not None:
        save_kwargs["dpi"] = (dpi, dpi)

    if image_format in {"JPEG", "BMP"} and image.mode not in {"RGB", "L"}:
        image = image.convert("RGB")

    image.save(destination, format=image_format, **save_kwargs)


def images_to_pdf(
    files: list[Path],
    output_dir: Path,
    progress_callback: ProgressCallback = None,
) -> ProcessingResult:
    output_dir.mkdir(parents=True, exist_ok=True)
    images: list[Image.Image] = []
    total_files = len(files)
    emit_progress(progress_callback, 0.02, "Preparing images")

    for file_index, file_path in enumerate(files, start=1):
        require_suffix(file_path, IMAGE_SUFFIXES, "image")
        image = Image.open(file_path)
        image = ImageOps.exif_transpose(image).convert("RGB")
        images.append(image)
        emit_progress(
            progress_callback,
            overall_ratio(file_index, total_files, 0.75),
            f"Loading image {file_index} of {total_files}",
        )

    if not images:
        raise ProcessingError("Upload at least one image file.")

    output_path = output_dir / "images-to-pdf.pdf"
    first_image, *remaining = images
    emit_progress(progress_callback, 0.9, "Saving combined PDF")
    first_image.save(output_path, save_all=True, append_images=remaining)

    for image in images:
        image.close()

    emit_progress(progress_callback, 1.0, "Image PDF ready")
    return ProcessingResult(files=[output_path], message=f"Combined {len(images)} image(s) into a PDF.")


def pdf_to_images(
    files: list[Path],
    output_dir: Path,
    progress_callback: ProgressCallback = None,
) -> ProcessingResult:
    output_dir.mkdir(parents=True, exist_ok=True)
    source = files[0]
    require_suffix(source, {".pdf"}, "PDF")

    document = fitz.open(source)
    total_pages = document.page_count or 1
    exported: list[Path] = []
    emit_progress(progress_callback, 0.02, "Reading PDF pages")

    for index, page in enumerate(document, start=1):
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
        target = output_dir / f"{source.stem}-page-{index:03d}.jpg"
        image = Image.open(io.BytesIO(pixmap.tobytes("png"))).convert("RGB")
        image.save(target, "JPEG", quality=92)
        image.close()
        exported.append(target)
        emit_progress(
            progress_callback,
            index / total_pages,
            f"Rendering page {index} of {total_pages}",
        )
    document.close()

    return ProcessingResult(files=exported, message=f"Exported {len(exported)} JPG page(s).")


def resize_and_convert_images(
    files: list[Path],
    output_dir: Path,
    width_px,
    height_px,
    dpi,
    output_format: str,
    keep_aspect: bool,
    progress_callback: ProgressCallback = None,
) -> ProcessingResult:
    output_dir.mkdir(parents=True, exist_ok=True)
    width = parse_optional_positive_int(width_px, "Width")
    height = parse_optional_positive_int(height_px, "Height")
    dpi_value = parse_optional_positive_int(dpi, "DPI")
    output_key = clean_text(output_format).lower() or "original"
    if output_key not in IMAGE_EXPORT_FORMATS:
        raise ProcessingError("Choose a valid output image format.")

    if width is None and height is None and dpi_value is None and output_key == "original":
        raise ProcessingError("Enter a width, height, DPI, or choose a different output format.")

    exported: list[Path] = []
    total_files = len(files)
    resampling = getattr(Image, "Resampling", Image).LANCZOS
    emit_progress(progress_callback, 0.02, "Preparing image conversion")

    for file_index, file_path in enumerate(files, start=1):
        require_suffix(file_path, IMAGE_SUFFIXES, "image")
        target_format, target_suffix = IMAGE_EXPORT_FORMATS[output_key]
        if target_format is None or target_suffix is None:
            target_format, target_suffix = infer_source_image_format(file_path)

        with Image.open(file_path) as opened_image:
            image = ImageOps.exif_transpose(opened_image)
            source_width, source_height = image.size
            new_width, new_height = resolve_resize_dimensions(width, height, source_width, source_height, keep_aspect)

            if (new_width, new_height) != image.size:
                image = image.resize((new_width, new_height), resample=resampling)
            else:
                image = image.copy()

            output_path = output_dir / f"{file_path.stem}-{new_width}x{new_height}{target_suffix}"
            save_image(image, output_path, target_format, dpi_value)
            image.close()
            exported.append(output_path)

        emit_progress(
            progress_callback,
            file_index / total_files,
            f"Processing image {file_index} of {total_files}",
        )

    message_parts = [f"Processed {len(exported)} image(s)"]
    if width is not None or height is not None:
        message_parts.append("with pixel resizing")
    if dpi_value is not None:
        message_parts.append(f"at {dpi_value} DPI")
    if output_key != "original":
        message_parts.append(f"as {output_key.upper()}")
    return ProcessingResult(files=exported, message=" ".join(message_parts) + ".")


def word_to_pdf(
    files: list[Path],
    output_dir: Path,
    progress_callback: ProgressCallback = None,
) -> ProcessingResult:
    output_dir.mkdir(parents=True, exist_ok=True)
    exported: list[Path] = []
    styles = getSampleStyleSheet()
    total_files = len(files)
    emit_progress(progress_callback, 0.02, "Reading Word files")

    for file_index, file_path in enumerate(files, start=1):
        require_suffix(file_path, {".docx"}, "Word")
        document = Document(file_path)
        output_path = output_dir / f"{file_path.stem}.pdf"
        story = []

        if not document.paragraphs and not document.tables:
            raise ProcessingError(f"{file_path.name} does not contain readable content.")

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                story.append(Spacer(1, 8))
                continue
            style_name = paragraph.style.name.lower() if paragraph.style and paragraph.style.name else ""
            if "title" in style_name:
                style = styles["Title"]
            elif "heading" in style_name:
                style = styles["Heading2"]
            else:
                style = styles["BodyText"]
            story.append(Paragraph(html.escape(text), style))
            story.append(Spacer(1, 8))

        for table in document.tables:
            matrix = [[clean_text(cell.text) for cell in row.cells] for row in table.rows]
            if not matrix:
                continue
            pdf_table = Table(matrix, repeatRows=1)
            pdf_table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f2937")),
                        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cbd5e1")),
                        ("PADDING", (0, 0), (-1, -1), 6),
                    ]
                )
            )
            story.append(pdf_table)
            story.append(Spacer(1, 14))

        pdf = SimpleDocTemplate(str(output_path), pagesize=A4)
        pdf.build(story)
        exported.append(output_path)
        emit_progress(
            progress_callback,
            file_index / total_files,
            f"Converting Word file {file_index} of {total_files}",
        )

    return ProcessingResult(files=exported, message=f"Converted {len(exported)} Word file(s) to PDF.")


def extract_pdf_tables(page) -> list[list[list[str]]]:
    if not hasattr(page, "find_tables"):
        return []

    tables_result = page.find_tables()
    extracted = []
    for table in tables_result.tables:
        rows = table.extract()
        if rows:
            extracted.append([[clean_text(cell) for cell in row] for row in rows])
    return extracted


def pdf_to_word(
    files: list[Path],
    output_dir: Path,
    progress_callback: ProgressCallback = None,
) -> ProcessingResult:
    output_dir.mkdir(parents=True, exist_ok=True)
    exported: list[Path] = []
    total_files = len(files)
    emit_progress(progress_callback, 0.02, "Preparing PDF to Word conversion")

    for file_index, file_path in enumerate(files, start=1):
        require_suffix(file_path, {".pdf"}, "PDF")
        source_pdf = fitz.open(file_path)
        total_pages = source_pdf.page_count or 1
        target_doc = Document()
        target_doc.add_heading(file_path.stem, level=0)

        for index, page in enumerate(source_pdf, start=1):
            target_doc.add_heading(f"Page {index}", level=1)
            blocks = page.get_text("blocks")
            if not blocks:
                target_doc.add_paragraph("[No text detected]")
            else:
                ordered_blocks = sorted(blocks, key=lambda block: (block[1], block[0]))
                for block in ordered_blocks:
                    text = clean_text(block[4])
                    if text:
                        target_doc.add_paragraph(text)

            for table_rows in extract_pdf_tables(page):
                if not table_rows:
                    continue
                column_count = max(len(row) for row in table_rows)
                doc_table = target_doc.add_table(rows=len(table_rows), cols=column_count)
                doc_table.style = "Table Grid"
                for row_index, row in enumerate(table_rows):
                    for column_index, value in enumerate(row):
                        doc_table.cell(row_index, column_index).text = value

            if index != total_pages:
                target_doc.add_page_break()

            emit_progress(
                progress_callback,
                overall_ratio(file_index, total_files, index / total_pages),
                f"Extracting page {index} of {total_pages} from {file_path.name}",
            )

        output_path = output_dir / f"{file_path.stem}.docx"
        target_doc.save(output_path)
        source_pdf.close()
        exported.append(output_path)

    return ProcessingResult(files=exported, message=f"Created {len(exported)} DOCX file(s).")


def worksheet_to_matrix(sheet) -> list[list[str]]:
    rows = list(sheet.iter_rows(values_only=True))
    if not rows:
        return [["Sheet is empty"]]

    last_row_index = 0
    last_column_index = 0
    for row_index, row in enumerate(rows, start=1):
        if any(value not in (None, "") for value in row):
            last_row_index = row_index
            for column_index, value in enumerate(row, start=1):
                if value not in (None, ""):
                    last_column_index = max(last_column_index, column_index)

    if last_row_index == 0 or last_column_index == 0:
        return [["Sheet is empty"]]

    trimmed = rows[:last_row_index]
    matrix = []
    for row in trimmed:
        matrix.append([clean_text(value) for value in row[:last_column_index]])
    return matrix


def excel_to_pdf(
    files: list[Path],
    output_dir: Path,
    progress_callback: ProgressCallback = None,
) -> ProcessingResult:
    output_dir.mkdir(parents=True, exist_ok=True)
    exported: list[Path] = []
    styles = getSampleStyleSheet()
    total_files = len(files)
    emit_progress(progress_callback, 0.02, "Preparing Excel conversion")

    for file_index, file_path in enumerate(files, start=1):
        require_suffix(file_path, {".xlsx", ".xlsm"}, "Excel")
        workbook = load_workbook(file_path, data_only=True)
        output_path = output_dir / f"{file_path.stem}.pdf"
        story = []
        total_sheets = len(workbook.worksheets) or 1

        for sheet_index, sheet in enumerate(workbook.worksheets, start=1):
            story.append(Paragraph(html.escape(sheet.title), styles["Heading1"]))
            story.append(Spacer(1, 10))

            matrix = worksheet_to_matrix(sheet)
            column_chunks = chunk_list(list(range(len(matrix[0]))), 8)
            for chunk in column_chunks:
                chunked_rows = [[row[column_index] for column_index in chunk] for row in matrix]
                heading = f"Columns {get_column_letter(chunk[0] + 1)}-{get_column_letter(chunk[-1] + 1)}"
                story.append(Paragraph(heading, styles["Heading3"]))
                table = Table(chunked_rows, repeatRows=1)
                table.setStyle(
                    TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0f766e")),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#94a3b8")),
                            ("FONTSIZE", (0, 0), (-1, -1), 7),
                            ("LEADING", (0, 0), (-1, -1), 8),
                            ("PADDING", (0, 0), (-1, -1), 4),
                        ]
                    )
                )
                story.append(table)
                story.append(Spacer(1, 14))

            if sheet_index != total_sheets:
                story.append(PageBreak())

            emit_progress(
                progress_callback,
                overall_ratio(file_index, total_files, sheet_index / total_sheets),
                f"Formatting sheet {sheet_index} of {total_sheets} from {file_path.name}",
            )

        pdf = SimpleDocTemplate(str(output_path), pagesize=landscape(A4))
        pdf.build(story)
        workbook.close()
        exported.append(output_path)

    return ProcessingResult(files=exported, message=f"Converted {len(exported)} Excel file(s) to PDF.")


def pdf_to_excel(
    files: list[Path],
    output_dir: Path,
    progress_callback: ProgressCallback = None,
) -> ProcessingResult:
    output_dir.mkdir(parents=True, exist_ok=True)
    exported: list[Path] = []
    total_files = len(files)
    emit_progress(progress_callback, 0.02, "Preparing PDF to Excel conversion")

    for file_index, file_path in enumerate(files, start=1):
        require_suffix(file_path, {".pdf"}, "PDF")
        pdf_document = fitz.open(file_path)
        total_pages = pdf_document.page_count or 1
        workbook = Workbook()
        first_sheet = workbook.active
        workbook.remove(first_sheet)

        for page_index, page in enumerate(pdf_document, start=1):
            tables = extract_pdf_tables(page)
            sheet = workbook.create_sheet(title=f"Page {page_index}"[:31])
            current_row = 1

            if tables:
                for table_index, table_rows in enumerate(tables, start=1):
                    sheet.cell(row=current_row, column=1, value=f"Table {table_index}")
                    current_row += 1
                    for row in table_rows:
                        for column_index, value in enumerate(row, start=1):
                            sheet.cell(row=current_row, column=column_index, value=value)
                        current_row += 1
                    current_row += 1
            else:
                text_lines = [line for line in page.get_text("text").splitlines() if line.strip()]
                if not text_lines:
                    sheet.cell(row=1, column=1, value="[No text detected]")
                else:
                    for row_index, line in enumerate(text_lines, start=1):
                        parts = [part for part in re.split(r"\s{2,}|\t", line) if part]
                        if not parts:
                            parts = [line]
                        for column_index, value in enumerate(parts, start=1):
                            sheet.cell(row=row_index, column=column_index, value=value)

            emit_progress(
                progress_callback,
                overall_ratio(file_index, total_files, page_index / total_pages),
                f"Extracting page {page_index} of {total_pages} from {file_path.name}",
            )

        output_path = output_dir / f"{file_path.stem}.xlsx"
        workbook.save(output_path)
        pdf_document.close()
        exported.append(output_path)

    return ProcessingResult(files=exported, message=f"Created {len(exported)} XLSX file(s).")


def merge_pdfs(
    files: list[Path],
    output_dir: Path,
    progress_callback: ProgressCallback = None,
) -> ProcessingResult:
    output_dir.mkdir(parents=True, exist_ok=True)
    if len(files) < 2:
        raise ProcessingError("Upload at least two PDF files to merge.")

    writer = PdfWriter()
    total_files = len(files)
    emit_progress(progress_callback, 0.02, "Preparing merge")

    for file_index, file_path in enumerate(files, start=1):
        require_suffix(file_path, {".pdf"}, "PDF")
        writer.append(str(file_path))
        emit_progress(
            progress_callback,
            file_index / total_files,
            f"Merging PDF {file_index} of {total_files}",
        )

    output_path = output_dir / "merged.pdf"
    with output_path.open("wb") as handle:
        writer.write(handle)
    writer.close()
    return ProcessingResult(files=[output_path], message=f"Merged {len(files)} PDFs into one file.")


def split_pdf(
    files: list[Path],
    output_dir: Path,
    mode: str,
    ranges: str,
    progress_callback: ProgressCallback = None,
) -> ProcessingResult:
    output_dir.mkdir(parents=True, exist_ok=True)
    source = files[0]
    require_suffix(source, {".pdf"}, "PDF")
    reader = PdfReader(str(source))
    exported: list[Path] = []
    emit_progress(progress_callback, 0.02, "Preparing split")

    if mode == "ranges":
        page_ranges = parse_page_ranges(ranges, len(reader.pages))
        total_parts = len(page_ranges)
        for index, (start, end) in enumerate(page_ranges, start=1):
            writer = PdfWriter()
            for page_number in range(start - 1, end):
                writer.add_page(reader.pages[page_number])
            output_path = output_dir / f"{source.stem}-part-{index:02d}-{start}-{end}.pdf"
            with output_path.open("wb") as handle:
                writer.write(handle)
            exported.append(output_path)
            emit_progress(
                progress_callback,
                index / total_parts,
                f"Creating split part {index} of {total_parts}",
            )
        return ProcessingResult(files=exported, message=f"Created {len(exported)} split PDF part(s).")

    total_pages = len(reader.pages)
    for page_number, page in enumerate(reader.pages, start=1):
        writer = PdfWriter()
        writer.add_page(page)
        output_path = output_dir / f"{source.stem}-page-{page_number:03d}.pdf"
        with output_path.open("wb") as handle:
            writer.write(handle)
        exported.append(output_path)
        emit_progress(
            progress_callback,
            page_number / total_pages,
            f"Creating page {page_number} of {total_pages}",
        )

    return ProcessingResult(files=exported, message=f"Created {len(exported)} split PDF page(s).")


def remove_watermark(
    files: list[Path],
    output_dir: Path,
    watermark_text: str,
    progress_callback: ProgressCallback = None,
) -> ProcessingResult:
    output_dir.mkdir(parents=True, exist_ok=True)
    exported: list[Path] = []
    normalized_text = watermark_text.strip()
    total_files = len(files)
    emit_progress(progress_callback, 0.02, "Preparing watermark cleanup")

    for file_index, file_path in enumerate(files, start=1):
        require_suffix(file_path, {".pdf"}, "PDF")
        document = fitz.open(file_path)
        total_pages = document.page_count or 1
        for page_index, page in enumerate(document, start=1):
            annotations = list(page.annots() or [])
            for annotation in annotations:
                page.delete_annot(annotation)

            if normalized_text:
                rects = page.search_for(normalized_text)
                for rect in rects:
                    page.add_redact_annot(rect, fill=(1, 1, 1))
                if rects:
                    page.apply_redactions()

            inner_ratio = page_index / total_pages
            emit_progress(
                progress_callback,
                overall_ratio(file_index, total_files, inner_ratio),
                f"Cleaning page {page_index} of {total_pages} in {file_path.name}",
            )

        output_path = output_dir / f"{file_path.stem}-clean.pdf"
        document.save(output_path, garbage=4, deflate=True)
        document.close()
        exported.append(output_path)

    return ProcessingResult(files=exported, message=f"Cleaned {len(exported)} PDF file(s).")


def render_pdf_page_to_image(page) -> Image.Image:
    pixmap = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
    image = Image.open(io.BytesIO(pixmap.tobytes("png")))
    image.load()
    return image


def ensure_tesseract() -> None:
    candidates: list[str] = []
    configured = os.environ.get("TESSERACT_CMD")
    if configured:
        candidates.append(configured)

    source_root = Path(__file__).resolve().parents[1]
    executable_root = Path(sys.executable).resolve().parent if getattr(sys, "frozen", False) else source_root
    meipass_root = Path(getattr(sys, "_MEIPASS")).resolve() if getattr(sys, "frozen", False) and hasattr(sys, "_MEIPASS") else None

    search_roots = [root for root in [meipass_root, executable_root, source_root] if root is not None]
    for root in search_roots:
        candidates.extend(
            [
                str(root / "tesseract" / "tesseract.exe"),
                str(root / "tesseract.exe"),
                str(root / "tesseract" / "bin" / "tesseract.exe"),
                str(root / "tesseract" / "tesseract"),
                str(root / "tesseract"),
            ]
        )

    candidates.extend(
        [
            "/opt/homebrew/bin/tesseract",
            "/usr/local/bin/tesseract",
            "/usr/bin/tesseract",
        ]
    )

    for candidate in candidates:
        if candidate and Path(candidate).exists():
            pytesseract.pytesseract.tesseract_cmd = candidate
            return

    discovered = shutil.which("tesseract")
    if discovered:
        pytesseract.pytesseract.tesseract_cmd = discovered
        return

    raise ProcessingError(
        "OCR needs Tesseract. Install it on the machine or use the bundled OCR runtime from the packaged app."
    )


def ocr_image_to_pdf(source_image: Image.Image, destination: Path) -> None:
    pdf_bytes = pytesseract.image_to_pdf_or_hocr(source_image, extension="pdf")
    destination.write_bytes(pdf_bytes)


def estimate_ocr_units(files: list[Path]) -> int:
    units = 0
    for file_path in files:
        suffix = file_path.suffix.lower()
        if suffix in IMAGE_SUFFIXES:
            units += 1
        elif suffix == ".pdf":
            units += count_pdf_pages(file_path)
    return max(units, 1)


def ocr_files(
    files: list[Path],
    output_dir: Path,
    progress_callback: ProgressCallback = None,
) -> ProcessingResult:
    output_dir.mkdir(parents=True, exist_ok=True)
    ensure_tesseract()
    exported: list[Path] = []
    total_units = estimate_ocr_units(files)
    completed_units = 0
    emit_progress(progress_callback, 0.02, "Preparing OCR")

    for file_path in files:
        suffix = file_path.suffix.lower()
        if suffix in IMAGE_SUFFIXES:
            image = Image.open(file_path)
            image = ImageOps.exif_transpose(image).convert("RGB")
            output_path = output_dir / f"{file_path.stem}-ocr.pdf"
            ocr_image_to_pdf(image, output_path)
            image.close()
            exported.append(output_path)
            completed_units += 1
            emit_progress(
                progress_callback,
                completed_units / total_units,
                f"Running OCR on {file_path.name}",
            )
            continue

        if suffix != ".pdf":
            raise ProcessingError(f"{file_path.name} cannot be processed by OCR.")

        source_pdf = fitz.open(file_path)
        writer = PdfWriter()
        temp_files: list[Path] = []

        for index, page in enumerate(source_pdf, start=1):
            image = render_pdf_page_to_image(page)
            temp_pdf_path = output_dir / f"{file_path.stem}-ocr-page-{index:03d}.pdf"
            ocr_image_to_pdf(image, temp_pdf_path)
            temp_files.append(temp_pdf_path)
            image.close()
            completed_units += 1
            emit_progress(
                progress_callback,
                completed_units / total_units,
                f"OCR page {index} from {file_path.name}",
            )

        for temp_path in temp_files:
            writer.append(str(temp_path))

        output_path = output_dir / f"{file_path.stem}-ocr.pdf"
        with output_path.open("wb") as handle:
            writer.write(handle)
        writer.close()

        for temp_path in temp_files:
            temp_path.unlink(missing_ok=True)

        source_pdf.close()
        exported.append(output_path)

    return ProcessingResult(files=exported, message=f"OCR completed for {len(exported)} file(s).")
